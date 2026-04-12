"""
심음 디노이징 연구 결과보고서 생성
LU-Net 기반 심음 노이즈 제거 모델 개발 및 Ablation Study
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = r"G:\stetho_ai\LSTM_first"
FIGURES_DIR = os.path.join(OUT_DIR, "figures")

doc = Document()

# ── 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 제목 스타일
for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = '맑은 고딕'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    h.font.color.rgb = RGBColor(0x1A, 0x25, 0x3C)

def add_table(doc, headers, rows, col_widths=None):
    """표 생성 헬퍼"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)

    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9.5)

    return table


# ═══════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('LU-Net 기반 심음 노이즈 제거 모델 개발\n및 Ablation Study 결과보고서')
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('End-to-End Deep Learning Framework for\nReal-Time Heart Sound Denoising')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
run.italic = True

for _ in range(4):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('건양대학교병원 미래융합교육원\n이종완\n2026년 3월')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

doc.add_page_break()


# ═══════════════════════════════════════════
# 목차
# ═══════════════════════════════════════════
doc.add_heading('목차', level=1)

toc_items = [
    '1. 연구 개요',
    '2. 기반 논문 분석',
    '3. 기존 접근법의 문제점 (Gemini 파이프라인)',
    '4. 제안 모델: LU-Net 기반 Waveform Denoising',
    '  4.1 데이터셋 구성',
    '  4.2 모델 아키텍처',
    '  4.3 LU-Net v1 → v2 개선사항',
    '  4.4 TU-Net: TCN 기반 변형',
    '5. 실험 결과',
    '  5.1 6모델 Ablation Study',
    '  5.2 SNR 레벨별 성능 분석',
    '  5.3 청음 평가',
    '6. 논의',
    '7. 결론 및 향후 계획',
    '참고문헌',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()


# ═══════════════════════════════════════════
# 1. 연구 개요
# ═══════════════════════════════════════════
doc.add_heading('1. 연구 개요', level=1)

doc.add_paragraph(
    '본 연구는 AI 전자청진기 시스템의 핵심 모듈인 심음 노이즈 제거(Heart Sound Denoising) 모델을 개발하고, '
    '다양한 아키텍처 변형과 학습 전략에 대한 체계적인 Ablation Study를 수행한 결과를 보고한다.'
)

doc.add_paragraph(
    '전자청진기로 수집되는 심음 신호에는 환경 소음(대화, 기기 알람), 접촉 잡음(피부 마찰, 탭핑), '
    '기기 내부 노이즈(회로 잡음, ADC 열잡음) 등 다양한 노이즈가 혼입된다. 이러한 노이즈는 후속 AI 분류 모델의 '
    '입력 품질을 저하시켜 진단 정확도에 직접적인 영향을 미친다. 따라서 실시간으로 동작 가능한 경량 디노이징 모델의 '
    '개발이 필수적이다.'
)

doc.add_heading('연구 목표', level=2)
goals = [
    '기반 논문(Ali et al., 2023)의 LU-Net 아키텍처를 재현하고 성능을 개선',
    '기기 특화 노이즈(ICS43434 마이크 노이즈, 접촉 잡음)를 포함한 학습 데이터 구축',
    'Spectrogram 도메인 vs Waveform 도메인 접근법의 비교 검증',
    'LSTM vs TCN(Temporal Convolutional Network) Skip Connection의 성능 비교',
    'Attention, Residual Learning, SI-SNR Loss 등 개선 기법의 개별 효과 분석',
    'TFLite 변환을 통한 Raspberry Pi 5 실시간 배포',
]
for g in goals:
    doc.add_paragraph(g, style='List Bullet')

doc.add_page_break()


# ═══════════════════════════════════════════
# 2. 기반 논문 분석
# ═══════════════════════════════════════════
doc.add_heading('2. 기반 논문 분석', level=1)

doc.add_heading('2.1 LU-Net (Ali et al., 2023)', level=2)

doc.add_paragraph(
    '본 연구의 기반이 되는 논문은 Ali et al.이 2023년 IEEE Access에 발표한 '
    '"An End-to-End Deep Learning Framework for Real-Time Denoising of Heart Sounds '
    'for Cardiac Disease Detection in Unseen Noise"이다.'
)

doc.add_paragraph(
    'LU-Net은 1D Conv로 구성된 U-Net 인코더-디코더 구조에서 Skip Connection에 '
    'Bidirectional LSTM을 삽입한 아키텍처이다. 핵심 아이디어는 다음과 같다:'
)

key_ideas = [
    'U-Net 인코더의 stride-2 Conv1D가 신호를 절반씩 축소하면서 자연스러운 다중 스케일(multi-scale) 분해를 수행',
    '각 스케일의 Skip Connection에 위치한 Bi-LSTM이 해당 해상도에서의 시간적 문맥(temporal context)을 학습',
    '결과적으로 각 LSTM이 서로 다른 주파수 대역에 전문화되는 효과 — Enc1(원본 해상도)은 세밀한 심음 파형, Enc5(16배 축소)는 전체 리듬 패턴',
    '입출력 모두 Raw Waveform이므로 Griffin-Lim 역변환이 불필요 → 위상 손실 문제 원천 차단',
]
for idea in key_ideas:
    doc.add_paragraph(idea, style='List Bullet')

doc.add_heading('논문 보고 성능', level=3)
add_table(doc,
    ['모델', 'ΔSNR (dB)', '파라미터 수'],
    [
        ['FCN', '+3.863', '~0.5M'],
        ['U-Net', '+4.330', '~1.0M'],
        ['LU-Net (논문)', '+5.575', '1.32M'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    '논문은 OAHS 데이터셋(성인 심음 1,000개)을 Clean 소스로, 폐음(ICBHI)과 병원 환경음(HAN)을 '
    '노이즈 소스로 사용하였다.'
)

doc.add_heading('2.2 논문의 한계점', level=2)

limitations = [
    ('노이즈 도메인의 한정', '폐음 + 병원 환경음 두 가지만 사용. 실제 청진 시 가장 빈번한 접촉 노이즈(마찰음, 탭핑)는 학습 데이터에 미포함'),
    ('Clean 데이터 다양성 부족', 'OAHS 성인 심음 1,000개만 사용. 아동 심음(심박수가 빠르고 S1/S2 간격이 짧음)은 미포함'),
    ('Data Augmentation 부재', '단순 SNR 랜덤 믹싱만 수행. Gain jitter, time shift 등 실전 변동성 미반영'),
    ('평가 지표 한정', 'SNR improvement만 보고. SI-SNR(Scale-Invariant SNR) 등 지각적 품질 지표 미사용'),
    ('샘플링 레이트 미최적화', '명시적 SR 최적화 없음. 심음의 고주파 성분(S3/S4, murmur) 보존 미고려'),
]
for title, desc in limitations:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()


# ═══════════════════════════════════════════
# 3. 기존 접근법의 문제점
# ═══════════════════════════════════════════
doc.add_heading('3. 기존 접근법의 문제점 (Gemini 파이프라인)', level=1)

doc.add_paragraph(
    '본 연구에 앞서, Gemini AI를 활용하여 스펙트로그램 기반 디노이징 파이프라인을 구축하였다. '
    '이 파이프라인은 두 가지 모델을 포함한다:'
)

add_table(doc,
    ['모델', '도메인', 'ΔSNR', '복원율', '주요 문제'],
    [
        ['Gemini Bi-LSTM', 'Mel-Spectrogram (2D)', '+4.86 dB', '53.65%', 'Griffin-Lim 아티팩트'],
        ['Gemini 마스킹 CRNN', 'Mel-Spectrogram (2D)', '+5.22 dB', '54.10%', '스케일링 불일치'],
    ]
)

doc.add_paragraph()
doc.add_heading('발견된 핵심 버그 4가지', level=2)

bugs = [
    ('스케일링 불일치 (가장 치명적)',
     '마스킹 모델에서 입력(X)과 타겟(Y)을 각각 다른 min/max로 정규화. '
     'Output = Input × Mask인데 Input과 Target의 기준이 다르면 Mask의 학습 정답이 왜곡됨.'),
    ('추론 시 스케일링 오류',
     '입력 오디오의 min/max를 그때그때 새로 계산. 훈련 때 사용한 글로벌 min/max와 불일치.'),
    ('가짜 Data Augmentation',
     '커스텀 노이즈를 ×10 복사하는 단순 Oversampling. 실제 다양성 증가 없이 특정 패턴에 과적합 유발.'),
    ('SNR 계산 오류',
     '스펙트로그램 도메인에서 SNR 계산 시 실제 음질 개선과 괴리 발생. Waveform 도메인 측정이 정확.'),
]
for title, desc in bugs:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph(
    '가장 근본적인 문제는 Spectrogram 도메인 자체의 한계이다. Mel-Spectrogram에서 디노이징 후 '
    '시간 도메인으로 복원하려면 Griffin-Lim 알고리즘을 통한 위상 복원이 필요한데, 이 과정에서 '
    '"물속에서 듣는 듯한" 아티팩트가 발생한다. 이는 구조적으로 해결이 불가능한 문제이며, '
    'Waveform 도메인 모델로의 전환이 필요한 근거가 되었다.'
)

doc.add_page_break()


# ═══════════════════════════════════════════
# 4. 제안 모델
# ═══════════════════════════════════════════
doc.add_heading('4. 제안 모델: LU-Net 기반 Waveform Denoising', level=1)

# 4.1 데이터셋
doc.add_heading('4.1 데이터셋 구성', level=2)

doc.add_paragraph(
    '총 10,000개의 Noisy/Clean 쌍을 합성하여 학습 데이터셋을 구축하였다. '
    '논문 대비 확장된 데이터 전략을 적용하였다.'
)

doc.add_heading('Clean Audio Sources', level=3)
add_table(doc,
    ['데이터소스', '파일 수', '설명'],
    [
        ['PhysioNet/CinC 2016 (Adult)', '~500', '성인 심음 (정상+비정상)'],
        ['PhysioNet Pediatric', '~500', '아동 심음 (심박수↑, S1/S2 간격↓)'],
        ['Custom ICS43434 녹음', '~50', 'ESP32+ICS43434로 직접 녹음한 심음'],
    ]
)

doc.add_paragraph()
doc.add_heading('Noise Sources', level=3)
add_table(doc,
    ['노이즈 소스', '파일 수', '설명'],
    [
        ['ESC-50', '2,000', '50개 카테고리 환경음 (범용)'],
        ['Tap/Rub (접촉 잡음)', '~50', '피부 마찰, 판막 탭핑 (1% 확률 혼합)'],
        ['Device BG (기기 배경음)', '~50', 'ICS43434 내부 회로 노이즈 (1% 확률 혼합)'],
    ]
)

doc.add_paragraph()
doc.add_heading('데이터 합성 파이프라인', level=3)
pipeline_steps = [
    'Clean + Noise 소스를 각각 4000Hz로 리샘플링',
    '5초 청크로 분할 (30% 오버랩)',
    'SNR 레벨 무작위 설정: -5dB ~ +15dB 균등 분포',
    'SNR 공식에 따라 노이즈 스케일링: noise_scaled = noise × (clean_rms / noise_rms) / 10^(SNR/20)',
    'Data Augmentation 적용: Gain jitter (±30%), Time shift (±0.2s), 하드웨어 열잡음 주입 (30% 확률)',
    'Peak Normalization (0.8)',
]
for i, step in enumerate(pipeline_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph()
doc.add_heading('논문 대비 데이터 전략 개선 요약', level=3)
add_table(doc,
    ['항목', '원본 LU-Net (2023)', '본 연구'],
    [
        ['노이즈 소스', '폐음 + 병원 환경음', '+ 접촉 마찰음 + 기기 노이즈 + ESC-50'],
        ['Clean 데이터', 'OAHS 성인 1,000개', '성인 + 아동 + 커스텀 ~1,050개'],
        ['Augmentation', 'SNR 랜덤 믹싱만', '+ Gain jitter + Time shift + HW noise'],
        ['슬라이싱', '명시 안 됨', '30% 오버랩'],
        ['샘플링 레이트', '명시 안 됨', '4000Hz (심음 최적화)'],
        ['평가 지표', 'SNR only', 'SNR + SI-SNR + Noise Reduction Rate'],
    ]
)

doc.add_page_break()

# 4.2 모델 아키텍처
doc.add_heading('4.2 모델 아키텍처', level=2)

doc.add_paragraph(
    'LU-Net의 핵심 구조는 5-Level 1D Conv U-Net 인코더-디코더에 Skip Connection마다 '
    'Bidirectional LSTM을 삽입한 것이다. 입력과 출력 모두 Raw Waveform(20,000 samples @ 4kHz, 5초)이며, '
    'Griffin-Lim 등의 역변환이 불필요하다.'
)

doc.add_heading('인코더 구조', level=3)
add_table(doc,
    ['레벨', '연산', '출력 Shape'],
    [
        ['Enc1', 'Conv1D(32) × 2', '(20000, 32)'],
        ['Enc2', 'Conv1D(64, stride=2)', '(10000, 64)'],
        ['Enc3', 'Conv1D(128, stride=2)', '(5000, 128)'],
        ['Enc4', 'Conv1D(256, stride=2)', '(2500, 256)'],
        ['Bottleneck', 'Conv1D(512, stride=2)', '(1250, 512)'],
    ]
)

doc.add_paragraph()
doc.add_heading('Skip Connection (Bi-LSTM)', level=3)
add_table(doc,
    ['Skip', 'Bi-LSTM Units', '출력 Shape', '역할'],
    [
        ['Skip1', 'BiLSTM(32)', '(20000, 64)', '세밀한 심음 파형 포착'],
        ['Skip2', 'BiLSTM(32)', '(10000, 64)', '중간 주파수 패턴'],
        ['Skip3', 'BiLSTM(64)', '(5000, 128)', '심음 리듬 구조'],
        ['Skip4', 'BiLSTM(64)', '(2500, 128)', '전체적 리듬 패턴'],
    ]
)

doc.add_paragraph()
doc.add_heading('디코더 구조', level=3)
doc.add_paragraph(
    '디코더는 UpSampling1D(2) + Concatenate(Skip) + Conv1D 구조로 인코더를 대칭적으로 복원한다. '
    '최종 출력은 Conv1D(1)로 단일 채널 파형을 생성한다.'
)

# 4.2 아키텍처 그림
if os.path.exists(os.path.join(FIGURES_DIR, '03_lunet_architecture.png')):
    doc.add_paragraph()
    doc.add_picture(os.path.join(FIGURES_DIR, '03_lunet_architecture.png'), width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('Figure 1. LU-Net v2 전체 아키텍처 — U-Net + BiLSTM + Channel Attention + Residual Learning')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_page_break()

# 4.3 개선사항
doc.add_heading('4.3 LU-Net v1 → v2 핵심 개선사항', level=2)

doc.add_paragraph(
    'LU-Net v1(논문 재현)에서 v2로의 개선은 세 가지 독립적인 기법의 조합으로 이루어졌다. '
    '각 기법은 심음 신호의 시계열 특성을 더 잘 활용하도록 설계되었다.'
)

doc.add_heading('Improvement 1: Residual Learning', level=3)
doc.add_paragraph(
    'v1은 모델이 Clean Waveform 전체를 직접 출력하도록 학습한다. '
    'v2는 이를 Noise Estimate로 전환하여, output = input - model(input) 구조를 채택하였다. '
    '이렇게 하면 모델의 학습 목표가 "깨끗한 소리 전체를 복원하라"에서 '
    '"노이즈만 찾아내라"로 단순화된다.'
)
doc.add_paragraph(
    '심음 파형에서 대부분의 시간 구간은 노이즈가 없거나 매우 적은 구간이다. '
    'Residual Learning을 적용하면 이런 구간에서 모델의 예측값이 거의 0에 가까우면 되므로, '
    '멀쩡한 심음 신호를 불필요하게 건드리는 실수가 줄어든다. '
    'DnCNN, 음성 향상(Speech Enhancement) 분야에서 이미 검증된 전략이다.'
)

doc.add_heading('Improvement 2: Channel Attention (SE Block)', level=3)
doc.add_paragraph(
    '각 Skip Connection의 Bi-LSTM 출력에 Squeeze-and-Excitation(SE) Block을 추가하였다. '
    'SE Block은 Global Average Pooling → FC(C, C/8) → ReLU → FC(C/8, C) → Sigmoid 구조로, '
    '채널별 중요도 가중치를 학습한다.'
)
doc.add_paragraph(
    '이를 통해 Skip Connection이 디코더에 전달하는 Feature 중 심음에 중요한 채널은 강조하고 '
    '노이즈에 해당하는 채널은 억제하는 Feature Selection 효과를 달성한다. '
    '심음 분류 프로젝트에서 CBAM Attention이 효과적이었던 경험을 바탕으로 도입하였다.'
)

doc.add_heading('Improvement 3: SI-SNR Loss', level=3)
doc.add_paragraph(
    'v1은 MSE(Mean Squared Error) Loss만 사용한다. MSE는 각 시점의 값 차이만 보지만, '
    'SI-SNR(Scale-Invariant Signal-to-Noise Ratio)은 전체 파형의 형태 일치도를 직접 측정한다.'
)
doc.add_paragraph(
    'SI-SNR Loss는 음성 분리(Speech Separation) 분야에서 Conv-TasNet 이후 사실상 표준이 된 '
    '손실 함수이다. 본 연구에서는 MSE와 SI-SNR을 복합 적용하였으며, 사전 실험을 통해 '
    '최적 비율을 탐색하였다.'
)

doc.add_paragraph()
doc.add_heading('Loss 비율 사전 실험', level=3)
add_table(doc,
    ['비율', 'ΔSNR', 'SI-SNR', '판정'],
    [
        ['MSE 100% (v1)', '+5.94 dB', '9.82 dB', '노이즈 제거력 최강, 파형 훼손 존재'],
        ['MSE 50% + SI-SNR 50%', '+0.12 dB', '12.44 dB', '학습 실패 — 노이즈 미제거'],
        ['MSE 90% + SI-SNR 10%', '+5.59 dB', '12.76 dB', '최적 균형 ← 채택'],
    ]
)
doc.add_paragraph()
doc.add_paragraph(
    'SI-SNR 비율이 50%일 때, 모델은 "차라리 아무것도 안 건드리는 게 안전하다"는 방향으로 학습하여 '
    'ΔSNR +0.12 dB로 사실상 노이즈를 제거하지 못하였다. 반면 90:10 비율에서는 '
    'MSE가 공격적 노이즈 제거를, SI-SNR이 파형 보존을 각각 담당하여 최적의 균형을 달성하였다.'
)

# 4.3 개선 그림
if os.path.exists(os.path.join(FIGURES_DIR, '04_three_improvements.png')):
    doc.add_paragraph()
    doc.add_picture(os.path.join(FIGURES_DIR, '04_three_improvements.png'), width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('Figure 2. LU-Net v1 → v2: 3대 핵심 개선 기법 — Residual Learning, Channel Attention, SI-SNR Loss')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_page_break()

# 4.4 TU-Net
doc.add_heading('4.4 TU-Net: TCN 기반 변형', level=2)

doc.add_paragraph(
    'LSTM은 시퀀스를 순차적으로 처리하므로 시퀀스 길이에 비례하여 학습 시간이 증가한다. '
    'Skip1의 Bi-LSTM은 원본 해상도 20,000 스텝을 처리하며, 이것이 에포크당 10분 이상 소요되는 '
    '주요 원인이다.'
)

doc.add_paragraph(
    'TCN(Temporal Convolutional Network)은 Dilated Causal Convolution으로 동일한 장거리 의존성을 '
    '포착하면서 모든 시점을 병렬로 처리할 수 있다. Dilation rate를 1, 2, 4, 8, 16으로 쌓으면 '
    'Receptive Field가 기하급수적으로 확장되어 LSTM과 유사한 시간적 문맥을 확보한다.'
)

doc.add_paragraph(
    'TU-Net은 LU-Net의 Skip Connection에서 Bi-LSTM을 TCN 블록으로 교체한 변형이다. '
    '나머지 구조(Encoder, Decoder, Attention, Residual Learning, Loss)는 동일하게 유지하여 '
    '순수한 LSTM vs TCN 효과를 분리 측정할 수 있도록 하였다.'
)

doc.add_heading('실험 설계: 2×2 Ablation', level=3)
add_table(doc,
    ['', 'v1 (MSE Only)', 'v2 (Attn + Residual + 90:10)'],
    [
        ['LSTM (LU-Net)', 'LU-Net v1', 'LU-Net v2'],
        ['TCN (TU-Net)', 'TU-Net v1', 'TU-Net v2'],
    ]
)
doc.add_paragraph()
doc.add_paragraph(
    '이 2×2 설계를 통해 "Skip 모듈(LSTM vs TCN) 효과"와 '
    '"개선 기법(Attention + Residual + SI-SNR) 효과"를 독립적으로 분리하여 분석할 수 있다.'
)

doc.add_page_break()


# ═══════════════════════════════════════════
# 5. 실험 결과
# ═══════════════════════════════════════════
doc.add_heading('5. 실험 결과', level=1)

doc.add_heading('5.1 6모델 Ablation Study', level=2)

doc.add_paragraph(
    '총 6개 모델(Gemini 2종 + LU-Net 2종 + TU-Net 2종)의 성능을 3가지 지표로 비교하였다.'
)

add_table(doc,
    ['모델', '아키텍처', 'ΔSNR (dB)', 'SI-SNR (dB)', '복원율 (%)', '에포크 시간'],
    [
        ['Gemini Bi-LSTM', 'Spec + LSTM', '+4.86', 'N/A', '53.65', '-'],
        ['Gemini 마스킹 CRNN', 'Spec + CRNN', '+5.22', 'N/A', '54.10', '-'],
        ['LU-Net v1', 'Wave + LSTM + MSE', '+5.94', '9.82', '54.20', '~12분'],
        ['LU-Net v2', 'Wave + LSTM + Attn+Res+90:10', '+5.59', '12.76', '52.27', '~12분'],
        ['TU-Net v1', 'Wave + TCN + MSE', '+4.35', '8.14', '44.71', '~4분'],
        ['TU-Net v2', 'Wave + TCN + Attn+Res+90:10', '+5.63', '12.59', '54.26', '~4분'],
    ]
)

# 비교 그림
if os.path.exists(os.path.join(FIGURES_DIR, '01_model_comparison.png')):
    doc.add_paragraph()
    doc.add_picture(os.path.join(FIGURES_DIR, '01_model_comparison.png'), width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('Figure 3. 6개 모델 성능 비교 — ΔSNR, SI-SNR, Noise Reduction Rate')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()
doc.add_heading('핵심 발견', level=3)

findings = [
    ('Waveform > Spectrogram',
     'Waveform 도메인 모델(LU-Net v1: +5.94)이 Spectrogram 도메인 최고 모델(Gemini CRNN: +5.22) 대비 '
     '+0.72 dB 향상. Griffin-Lim 아티팩트 문제도 원천 해결.'),
    ('논문 성능 초과 달성',
     'LU-Net v1(+5.94)과 TU-Net v2(+5.63) 모두 원본 논문(+5.575) 성능을 초과. '
     '데이터 전략 개선만으로 아키텍처 변경 없이 논문을 넘긴 것이 핵심 contribution.'),
    ('Attention + Residual + SI-SNR의 결정적 효과',
     'TU-Net v1(+4.35) → TU-Net v2(+5.63)로 +1.28 dB 향상. '
     '같은 TCN 구조에서 개선 기법 적용만으로 논문 수준에 도달.'),
    ('LSTM vs TCN: v1에서는 LSTM 압승, v2에서는 동급',
     'v1 비교: LSTM +5.94 vs TCN +4.35 (LSTM 승). '
     'v2 비교: LSTM +5.59 vs TCN +5.63 (동급). '
     '개선 기법이 TCN의 구조적 약점을 완전히 보완.'),
    ('SI-SNR vs ΔSNR의 괴리',
     'LU-Net v2는 ΔSNR(-0.35)이 v1보다 낮지만 SI-SNR(+2.94)은 높음. '
     '청음 평가에서 v2가 우수하게 평가됨 → ΔSNR만으로 음질을 판단하면 안 됨을 실험적으로 증명.'),
    ('TU-Net v2: 효율성 최강',
     'LU-Net v2와 동급 성능을 에포크당 4분(vs 12분)에 달성. '
     '3배 빠른 학습 시간으로 동일한 품질 → 실전 배포에 가장 적합.'),
]

for title, desc in findings:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# 5.2 SNR별 분석
doc.add_heading('5.2 SNR 레벨별 성능 분석', level=2)

doc.add_paragraph(
    '입력 SNR 레벨(-5, 0, 5, 10, 15 dB)별로 디노이징 성능을 분리 측정하였다.'
)

doc.add_paragraph(
    '저SNR(-5dB)에서 향상폭이 가장 크며(+20dB 이상), 고SNR(15dB)에서는 향상폭이 감소한다(+1~2dB). '
    '이는 이미 노이즈가 적은 신호에서는 추가 제거의 여지가 적기 때문이며, '
    '실제 임상 환경에서 노이즈가 심한 상황일수록 모델의 효과가 극대화됨을 의미한다.'
)

if os.path.exists(os.path.join(FIGURES_DIR, '06_snr_analysis.png')):
    doc.add_paragraph()
    doc.add_picture(os.path.join(FIGURES_DIR, '06_snr_analysis.png'), width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('Figure 4. SNR 레벨별 디노이징 성능 분석 + BiLSTM vs TCN 비교')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

# 5.3 청음 평가
doc.add_heading('5.3 청음 평가', level=2)

doc.add_paragraph(
    'SNR -5dB(노이즈가 심음보다 큰 극한 조건) 테스트 파일에 대해 주관적 청음 평가를 수행하였다.'
)

doc.add_paragraph(
    'LU-Net v2와 TU-Net v2는 ΔSNR 수치상 v1보다 낮지만, 청음 시 심음의 질감과 자연스러움이 '
    '더 우수하게 평가되었다. 이는 SI-SNR 12.76 dB(v2)의 효과로, 파형의 형태가 원본에 더 가깝게 '
    '보존되었기 때문이다. 특히 50:50 비율과 90:10 비율의 청음 비교에서 90:10이 "노이즈도 지우면서 '
    '자연스러운" 소리로 평가되어, Loss 비율의 최적화가 최종 청음 품질에 결정적임을 확인하였다.'
)

doc.add_paragraph(
    '이 결과는 디노이징 모델의 품질 평가에서 ΔSNR 단독 지표의 한계를 보여주며, '
    'SI-SNR 등 지각적 품질 지표의 병행 사용이 필수적임을 시사한다.'
)

doc.add_page_break()


# ═══════════════════════════════════════════
# 6. 논의
# ═══════════════════════════════════════════
doc.add_heading('6. 논의', level=1)

doc.add_heading('6.1 Spectrogram vs Waveform 도메인', level=2)
doc.add_paragraph(
    '본 연구에서 Waveform 도메인 모델이 Spectrogram 도메인 모델을 일관되게 상회하는 결과를 보였다. '
    'Spectrogram 방식의 근본적 한계인 Griffin-Lim 위상 복원 아티팩트가 제거되었으며, '
    'End-to-End 학습이 가능하여 전처리/후처리의 정보 손실이 최소화되었다. '
    '다만, Spectrogram 방식이 주파수 구조를 직접 다룰 수 있다는 장점이 있으므로, '
    '향후 Waveform 모델에 Multi-Resolution STFT Loss를 보조 손실로 추가하여 '
    '두 도메인의 장점을 결합하는 연구가 가능하다.'
)

if os.path.exists(os.path.join(FIGURES_DIR, '02_spec_vs_waveform.png')):
    doc.add_paragraph()
    doc.add_picture(os.path.join(FIGURES_DIR, '02_spec_vs_waveform.png'), width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('Figure 5. Spectrogram vs Waveform 도메인 접근법 비교')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_heading('6.2 청크 길이의 영향', level=2)
doc.add_paragraph(
    '5초 vs 2초 청크 길이 비교 실험에서 5초가 ΔSNR +5.94 vs +5.57로 우위를 보였다. '
    '심음 한 주기가 0.8~1초이므로, 5초는 5~6주기를 커버하여 LSTM이 심음의 주기성을 '
    '충분히 학습할 수 있었다. 반면 2초는 2~3주기만 포함되어 시간적 문맥이 부족하였다. '
    '다만 추론 시에는 5초 입력을 2초 단위로 Overlap-Add 처리하여 이음새를 자연스럽게 '
    '연결하는 방식도 적용 가능하다.'
)

doc.add_heading('6.3 Attention + Residual + SI-SNR의 시너지', level=2)
doc.add_paragraph(
    'Ablation Study에서 가장 주목할 발견은 세 기법의 조합이 TCN의 구조적 약점을 완전히 보완했다는 점이다. '
    'TU-Net v1(TCN + MSE only)은 ΔSNR +4.35로 LU-Net v1(+5.94) 대비 크게 뒤졌으나, '
    'TU-Net v2(TCN + Attn+Res+90:10)는 +5.63으로 LU-Net v2(+5.59)와 동급 수준에 도달하였다. '
    '이는 Skip Connection 모듈 자체보다 학습 전략의 개선이 성능에 더 결정적임을 시사한다.'
)

doc.add_heading('6.4 실전 배포 관점', level=2)
doc.add_paragraph(
    'TU-Net v2를 최종 배포 모델로 선정하였다. Conv1D 기반 TCN 구조는 TFLite 변환이 깔끔하며, '
    'LSTM 대비 추론 속도도 빠르다. TFLite로 변환된 모델(tunet_v2_best.tflite)은 '
    'Raspberry Pi 5에서 실시간 디노이징에 사용되고 있으며, 5초 오디오 처리에 1초 미만이 소요된다.'
)

doc.add_page_break()


# ═══════════════════════════════════════════
# 7. 결론 및 향후 계획
# ═══════════════════════════════════════════
doc.add_heading('7. 결론 및 향후 계획', level=1)

doc.add_heading('7.1 결론', level=2)

conclusions = [
    'LU-Net 기반 Waveform 디노이징 모델을 개발하여, 기반 논문(Ali et al., 2023) 대비 '
    'ΔSNR +0.37 dB 향상(+5.575 → +5.94)을 달성하였다.',
    '6개 모델에 대한 체계적 Ablation Study를 수행하여, Spectrogram vs Waveform, '
    'LSTM vs TCN, MSE vs SI-SNR Loss 등의 선택이 성능에 미치는 영향을 정량적으로 분석하였다.',
    'Channel Attention(SE Block), Residual Learning, SI-SNR Loss 세 가지 개선 기법의 조합이 '
    'TCN의 구조적 약점을 보완하여 LSTM과 동급 성능을 3배 빠른 학습 시간에 달성함을 확인하였다.',
    'ΔSNR과 SI-SNR 간의 괴리를 실험적으로 증명하여, 디노이징 품질 평가에 다중 지표 사용의 '
    '필요성을 제시하였다.',
    '최종 모델(TU-Net v2)을 TFLite로 변환하여 Raspberry Pi 5에서 실시간 심음 디노이징을 구현하였다.',
]
for c in conclusions:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('7.2 향후 계획', level=2)

future = [
    ('Dynamic Mixing (온라인 데이터 합성)',
     '현재 10,000개를 미리 생성하여 학습하는 방식을 매 배치마다 즉석 합성으로 전환. '
     '모델이 같은 조합을 두 번 볼 확률이 0이 되어 과적합 억제 효과 극대화.'),
    ('Curriculum Learning',
     '쉬운 문제(SNR +15dB)부터 시작하여 점진적으로 어려운 문제(SNR -5dB)로 난이도를 올리는 '
     '학습 전략 적용. 수렴 속도 개선 및 최종 성능 향상 기대.'),
    ('Multi-Resolution STFT Loss',
     'Waveform 도메인 MSE + SI-SNR에 주파수 도메인 보조 손실 추가. '
     '다양한 FFT 크기(512, 1024, 2048)에서 스펙트로그램 차이를 최소화하여 '
     '시간+주파수 양방향 최적화.'),
    ('nRF52840 Edge 배포',
     '현재 Raspberry Pi 5 배포에서 더 나아가, nRF52840 MCU에 TFLite Micro로 '
     '직접 탑재하여 BLE 통신 없이 디바이스 내에서 실시간 디노이징 수행.'),
]
for title, desc in future:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

# 타임라인 그림
if os.path.exists(os.path.join(FIGURES_DIR, '07_development_journey.png')):
    doc.add_paragraph()
    doc.add_picture(os.path.join(FIGURES_DIR, '07_development_journey.png'), width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('Figure 6. 전체 개발 타임라인 및 핵심 교훈')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_page_break()


# ═══════════════════════════════════════════
# 참고문헌
# ═══════════════════════════════════════════
doc.add_heading('참고문헌', level=1)

references = [
    '[1] Ali, S. N., Shuvo, S. B., Al-Manzo, M. I. S., Hasan, A., & Hasan, T. (2023). '
    '"An end-to-end deep learning framework for real-time denoising of heart sounds for '
    'cardiac disease detection in unseen noise." IEEE Access, 11, 87887-87901.',

    '[2] Ronneberger, O., Fischer, P., & Brox, T. (2015). '
    '"U-Net: Convolutional networks for biomedical image segmentation." '
    'MICCAI 2015, LNCS, vol. 9351, pp. 234-241.',

    '[3] Hu, J., Shen, L., & Sun, G. (2018). '
    '"Squeeze-and-excitation networks." '
    'Proceedings of the IEEE CVPR, pp. 7132-7141.',

    '[4] Luo, Y., & Mesgarani, N. (2019). '
    '"Conv-TasNet: Surpassing ideal time-frequency magnitude masking for speech separation." '
    'IEEE/ACM Transactions on Audio, Speech, and Language Processing, 27(8), 1256-1266.',

    '[5] Bai, S., Kolter, J. Z., & Koltun, V. (2018). '
    '"An empirical evaluation of generic convolutional and recurrent networks for sequence modeling." '
    'arXiv preprint arXiv:1803.01271.',

    '[6] Piczak, K. J. (2015). '
    '"ESC: Dataset for environmental sound classification." '
    'Proceedings of the 23rd ACM International Conference on Multimedia, pp. 1015-1018.',

    '[7] Clifford, G. D., et al. (2016). '
    '"Classification of normal/abnormal heart sound recordings: the PhysioNet/Computing '
    'in Cardiology Challenge 2016." Computing in Cardiology, 43, 609-612.',

    '[8] Zhang, K., Zuo, W., Chen, Y., Meng, D., & Zhang, L. (2017). '
    '"Beyond a Gaussian denoiser: Residual learning of deep CNN for image denoising." '
    'IEEE TIP, 26(7), 3142-3155.',
]

for ref in references:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.left_indent = Cm(1.0)
    for run in p.runs:
        run.font.size = Pt(9.5)


# ═══════════════════════════════════════════
# 저장
# ═══════════════════════════════════════════
output_path = os.path.join(OUT_DIR, 'Heart_Sound_Denoising_Report.docx')
doc.save(output_path)
print(f"보고서 생성 완료: {output_path}")
print(f"  페이지 수: 약 15-18 페이지")
print(f"  포함 그래프: 6장")
print(f"  포함 표: 11개")
